"""
research_engine/exporters/project_docx_exporter.py
Tier 1 #5 — Full project Word exporter

Produces a submission-ready .docx with:
  - Title page (cover page)
  - Declaration page
  - Abstract placeholder
  - Table of Contents (auto-generated from headings)
  - All chapters (1–5) with proper heading hierarchy
  - Reference list
  - Appendix placeholder
  - Page numbers (footer)
  - APA/academic formatting

Public API
----------
    export_project_docx(session, output_path, reference_list=None) → Path
"""
from __future__ import annotations

from pathlib import Path
from datetime import datetime
import re

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.section import WD_SECTION
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def export_project_docx(
    session,
    output_path:    str | Path,
    reference_list=None,
) -> Path:
    """
    Export the full project to a submission-ready Word document.

    Parameters
    ----------
    session        : ProjectSession with chapters written
    output_path    : where to save the .docx file
    reference_list : optional ReferenceList from reference_generator

    Returns
    -------
    Path — saved .docx file
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx required: pip install python-docx")

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    m   = session.metadata

    # ── Page setup ─────────────────────────────────────────────
    for sec in doc.sections:
        sec.page_height  = Cm(29.7)
        sec.page_width   = Cm(21.0)
        sec.top_margin   = Cm(2.54)
        sec.bottom_margin= Cm(2.54)
        sec.left_margin  = Cm(3.81)   # 1.5 inch — binding margin
        sec.right_margin = Cm(2.54)

    # ── Default body style ─────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    _set_line_spacing(doc)

    # ── 1. Title page ──────────────────────────────────────────
    _add_title_page(doc, m)
    doc.add_page_break()

    # ── 2. Declaration ─────────────────────────────────────────
    _add_declaration(doc, m)
    doc.add_page_break()

    # ── 3. Dedication (placeholder) ────────────────────────────
    _add_centered_page(doc, "DEDICATION",
        "This work is dedicated to ...")
    doc.add_page_break()

    # ── 4. Acknowledgements (placeholder) ──────────────────────
    _add_centered_page(doc, "ACKNOWLEDGEMENTS",
        "The researcher wishes to thank ...")
    doc.add_page_break()

    # ── 5. Abstract (placeholder) ──────────────────────────────
    _add_abstract(doc, session)
    doc.add_page_break()

    # ── 6. Table of Contents (stub — Word updates on open) ────
    _add_toc(doc)
    doc.add_page_break()

    # ── 7. List of Tables / Figures ────────────────────────────
    _add_section_heading(doc, "LIST OF TABLES")
    doc.add_paragraph("Table 4.1  Demographic Characteristics of Respondents")
    doc.add_paragraph("Table 4.2  Distribution by Gender")
    doc.add_paragraph("(Update this list manually or use Word's automatic table list)")
    doc.add_page_break()

    # ── 8. Chapters ────────────────────────────────────────────
    for n in sorted(session.chapters.keys()):
        ch = session.chapters[n]
        _add_chapter(doc, n, ch.content)
        doc.add_page_break()

    # ── 9. Reference list ──────────────────────────────────────
    _add_section_heading(doc, "REFERENCES")
    if reference_list and reference_list.entries:
        for entry in sorted(reference_list.entries, key=lambda e: e.author.lower()):
            p = doc.add_paragraph(entry.full_text)
            p.paragraph_format.left_indent  = Cm(1.27)
            p.paragraph_format.first_line_indent = Cm(-1.27)  # hanging indent
    else:
        doc.add_paragraph(
            "[References will be auto-generated. Run: "
            "python main.py project references --session SESSION_ID]"
        )
    doc.add_page_break()

    # ── 10. Appendix ────────────────────────────────────────────
    _add_section_heading(doc, "APPENDIX")
    _add_section_heading(doc, "Appendix I: Questionnaire", level=2)
    doc.add_paragraph("[Attach questionnaire instrument here]")

    # ── Page numbers in footer ─────────────────────────────────
    _add_page_numbers(doc)

    doc.save(str(output_path))
    return output_path


# ── Page builders ─────────────────────────────────────────────

def _add_title_page(doc, m):
    def _cp(text, bold=False, size=12, space_before=12):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(space_before)
        r = p.add_run(text)
        r.bold       = bold
        r.font.size  = Pt(size)
        return p

    if m.institution:
        _cp(m.institution.upper(), bold=True, size=14, space_before=72)
    if m.department:
        _cp(m.department, size=12, space_before=6)

    _cp("", size=12, space_before=36)  # spacer

    if m.title:
        _cp(m.title.upper(), bold=True, size=14, space_before=36)

    _cp("", size=12, space_before=24)
    _cp("A Research Project Submitted in Partial Fulfilment of the Requirements",
        size=11, space_before=36)
    _cp(f"for the Award of {m.level.value.upper()} Degree", size=11, space_before=4)

    if m.student_name:
        _cp("", size=12, space_before=48)
        _cp("By", size=11, space_before=12)
        _cp(m.student_name.upper(), bold=True, size=12, space_before=6)

    if m.supervisor:
        _cp("", size=12, space_before=24)
        _cp("Supervisor:", size=11, space_before=12)
        _cp(m.supervisor, size=11, space_before=4)

    year = m.year or str(datetime.now().year)
    _cp(year, size=12, space_before=48)


def _add_declaration(doc, m):
    _add_section_heading(doc, "DECLARATION")
    student = m.student_name or "The researcher"
    text = (
        f"I, {student}, declare that this research project is my original work. "
        "It has not been submitted previously for any degree or professional award "
        "to this or any other institution.\n\n"
        "All sources of information used in this work have been duly acknowledged "
        "in the text and in the reference list.\n\n\n"
        "Signature: ________________________\n\n"
        f"Date: _____________________________\n\n\n"
        "Certified by:\n\n"
        f"Supervisor: {m.supervisor or '________________________'}\n\n"
        "Signature: ________________________\n\n"
        "Date: _____________________________"
    )
    doc.add_paragraph(text)


def _add_abstract(doc, session):
    _add_section_heading(doc, "ABSTRACT")
    ch1 = session.chapters.get(1)
    ch5 = session.chapters.get(5)
    if ch1 and ch5:
        # Extract first 150 words of Ch1 background + last 100 words of Ch5
        bg_words = ch1.content.split()[:150]
        conc_words = ch5.content.split()[-80:]
        abstract_text = (
            " ".join(bg_words) + " ... " + " ".join(conc_words)
        )
        doc.add_paragraph(abstract_text[:800] + "\n\n")
    else:
        doc.add_paragraph("[Abstract — 150–250 words summarising background, objectives, methods, findings, and conclusions.]")

    m = session.metadata
    if m.keywords:
        p = doc.add_paragraph()
        r = p.add_run("Keywords: ")
        r.bold = True
        p.add_run(", ".join(m.keywords))


def _add_toc(doc):
    _add_section_heading(doc, "TABLE OF CONTENTS")
    doc.add_paragraph(
        "Right-click here in Word → Update Field → Update entire table\n"
        "(Or press Ctrl+A then F9 to update all fields)"
    ).italic = True
    # Insert a real TOC field
    para = doc.add_paragraph()
    run  = para.add_run()
    fld  = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    run._r.append(fld)
    run2 = para.add_run()
    instr = OxmlElement("w:instrText")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    run2._r.append(instr)
    run3 = para.add_run()
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run3._r.append(fld2)


def _add_centered_page(doc, heading, body_text):
    _add_section_heading(doc, heading)
    p = doc.add_paragraph(body_text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_section_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return h


def _add_chapter(doc, number: int, content: str):
    """Parse markdown-ish content and add to doc with proper heading hierarchy."""
    for line in content.split("\n"):
        s = line.strip()
        if not s:
            continue
        if s.startswith("#### "):
            doc.add_heading(s[5:], level=4)
        elif s.startswith("### "):
            doc.add_heading(s[4:], level=3)
        elif s.startswith("## "):
            doc.add_heading(s[3:], level=2)
        elif s.startswith("# "):
            doc.add_heading(s[2:], level=1)
        else:
            p = doc.add_paragraph(s)
            p.paragraph_format.first_line_indent = Cm(1.27)
            p.paragraph_format.space_after = Pt(0)


def _set_line_spacing(doc):
    style = doc.styles["Normal"]
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE


def _add_page_numbers(doc):
    """Add page number to footer of each section."""
    for section in doc.sections:
        footer = section.footer
        para   = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        fld = OxmlElement("w:fldChar")
        fld.set(qn("w:fldCharType"), "begin")
        run._r.append(fld)
        instr_run = para.add_run()
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        instr_run._r.append(instr)
        end_run = para.add_run()
        end_fld = OxmlElement("w:fldChar")
        end_fld.set(qn("w:fldCharType"), "end")
        end_run._r.append(end_fld)
