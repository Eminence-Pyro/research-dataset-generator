"""
research_engine/exporters/word_exporter.py
Stage 10 — Export Engine  |  Milestone 1.1.A

Produces a submission-ready Chapter Four Word document (.docx) containing:

    1. Cover block    — study title, date, seed
    2. Section tables — one descriptive statistics table per questionnaire section
                        (Item number | Statement | N | Mean | SD | Interpretation)
    3. Summary table  — overall section means and grand mean
    4. Frequency tables — one table per categorical demographic variable
    5. Crosstabulation tables — with χ², df, p-value, Cramer's V

Tables are formatted in the APA-adjacent style common in Nigerian academic
dissertations: bold headers, bordered cells, alternating row shading.

Public API
----------
    export_word(
        dataset, questionnaire, variable_dictionary,
        validation_report, likert_sum, freq_tables, crosstab_results,
        output_dir, study_title, seed
    ) → Path

Dependencies
------------
    python-docx >= 1.0  (pip install python-docx)
"""
from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from research_engine.models import Dataset, Questionnaire, VariableDictionary
from research_engine.validators.dataset_validator import ValidationReport
from research_engine.analysis.descriptives import LikertSummary
from research_engine.analysis.frequencies  import FrequencyTable
from research_engine.analysis.crosstabs    import CrosstabResult
from research_engine.analysis.reliability   import ReliabilityReport, SectionReliability


# ── Colour palette (RGB) ──────────────────────────────────────
_NAVY    = RGBColor(0x1F, 0x38, 0x64)   # header background
_WHITE   = RGBColor(0xFF, 0xFF, 0xFF)   # header text
_LTBLUE  = RGBColor(0xDC, 0xE6, 0xF1)  # odd data rows
_LTGREY  = RGBColor(0xF2, 0xF2, 0xF2)  # even data rows
_TOTAL   = RGBColor(0x17, 0x37, 0x5E)  # total row background
_BLACK   = RGBColor(0x00, 0x00, 0x00)


# ══════════════════════════════════════════════════════════════
# Low-level XML helpers
# ══════════════════════════════════════════════════════════════

def _set_cell_bg(cell, rgb: RGBColor) -> None:
    """Fill a table cell with a solid background colour."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    hex_colour = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_colour)
    tcPr.append(shd)


def _set_cell_border(cell) -> None:
    """Apply thin black borders to all four sides of a cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),   "single")
        b.set(qn("w:sz"),    "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "000000")
        borders.append(b)
    tcPr.append(borders)


def _set_col_width(table, col_idx: int, width_cm: float) -> None:
    """Set the width of a column in centimetres."""
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)


# ══════════════════════════════════════════════════════════════
# Document-level helpers
# ══════════════════════════════════════════════════════════════

def _set_margins(doc: Document,
                 top: float = 2.54, bottom: float = 2.54,
                 left: float = 3.17, right: float = 3.17) -> None:
    """Set page margins in centimetres (default: A4 thesis margins)."""
    section = doc.sections[0]
    section.top_margin    = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin   = Cm(left)
    section.right_margin  = Cm(right)


def _heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = _NAVY
    run.font.bold      = True


def _body(doc: Document, text: str, italic: bool = False) -> None:
    p   = doc.add_paragraph(text)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.size   = Pt(11)
    run.font.italic = italic


def _caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size   = Pt(11)
    run.font.bold   = True
    run.font.italic = False


# ══════════════════════════════════════════════════════════════
# Table builders
# ══════════════════════════════════════════════════════════════

def _make_table(doc: Document, rows: int, cols: int) -> Any:
    """Add a table with consistent style settings."""
    table = doc.add_table(rows=rows, cols=cols)
    table.style           = "Table Grid"
    table.alignment       = WD_TABLE_ALIGNMENT.CENTER
    return table


def _header_row(table, col_headers: list[str],
                widths: list[float] | None = None) -> None:
    """Write and format the header row of a table."""
    row = table.rows[0]
    for i, text in enumerate(col_headers):
        cell = row.cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(text)
        run.font.bold        = True
        run.font.size        = Pt(10)
        run.font.color.rgb   = _WHITE
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_bg(cell, _NAVY)
        _set_cell_border(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if widths:
        for i, w in enumerate(widths):
            if w:
                for row2 in table.rows:
                    row2.cells[i].width = Cm(w)


def _data_row(table, row_idx: int, values: list[Any],
              bold_first: bool = False, shade: bool = False,
              total: bool = False) -> None:
    """Write a data row with optional shading."""
    row  = table.rows[row_idx]
    fill = _TOTAL if total else (_LTBLUE if shade else _LTGREY)
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.text = ""
        run  = cell.paragraphs[0].add_run(str(val) if val is not None else "")
        run.font.size = Pt(10)
        run.font.bold = (total or (bold_first and i == 0))
        if total:
            run.font.color.rgb = _WHITE
        align = (WD_ALIGN_PARAGRAPH.LEFT
                 if i == 0 else WD_ALIGN_PARAGRAPH.CENTER)
        cell.paragraphs[0].alignment = align
        _set_cell_bg(cell, fill)
        _set_cell_border(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


# ══════════════════════════════════════════════════════════════
# Section tables — descriptive statistics per questionnaire section
# ══════════════════════════════════════════════════════════════

def _write_section_table(doc: Document, section_key: str,
                          section_title: str,
                          items: list,          # list[LikertItemStats]
                          section_mean: float,
                          table_num: int) -> None:
    """
    Write one Chapter Four descriptive statistics table for a questionnaire section.

    Columns: Item No. | Statement | N | Mean | SD | Interpretation
    Footer:  Section mean row
    """
    doc.add_paragraph()   # breathing space
    _caption(doc, f"Table {table_num}: {section_title}")
    doc.add_paragraph()

    n_rows = 1 + len(items) + 1   # header + items + section mean
    table  = _make_table(doc, n_rows, 6)
    headers = ["Item No.", "Statement", "N", "Mean", "SD", "Interpretation"]
    widths  = [1.8, 7.0, 1.0, 1.3, 1.2, 3.2]
    _header_row(table, headers, widths)

    for idx, item in enumerate(items):
        shade = (idx % 2 == 0)
        _data_row(
            table, idx + 1,
            [item.question_number, item.label, item.n,
             f"{item.mean:.2f}", f"{item.std:.2f}", item.interpretation],
            bold_first=False, shade=shade,
        )

    # Section mean footer row
    _data_row(
        table, len(items) + 1,
        [f"Section {section_key} Mean", "", items[0].n if items else "",
         f"{section_mean:.2f}", "", ""],
        total=True,
    )
    doc.add_paragraph()


# ══════════════════════════════════════════════════════════════
# Summary table — all section means + grand mean
# ══════════════════════════════════════════════════════════════

def _write_summary_table(doc: Document, likert_sum: LikertSummary,
                          questionnaire: Questionnaire,
                          table_num: int) -> None:
    """Grand summary: section means and overall mean."""
    doc.add_paragraph()
    _caption(doc, f"Table {table_num}: Summary of Mean Scores by Section")
    doc.add_paragraph()

    n_rows = 1 + len(questionnaire.sections) + 1
    table  = _make_table(doc, n_rows, 4)
    headers = ["Section", "Section Title", "Mean Score", "Interpretation"]
    widths  = [1.5, 8.5, 2.5, 3.0]
    _header_row(table, headers, widths)

    interp_fn = _interpret_mean

    for idx, sec in enumerate(questionnaire.sections):
        mean  = likert_sum.section_means.get(sec.key, 0.0)
        shade = (idx % 2 == 0)
        _data_row(
            table, idx + 1,
            [f"Section {sec.key}", sec.title, f"{mean:.2f}", interp_fn(mean)],
            shade=shade,
        )

    _data_row(
        table, len(questionnaire.sections) + 1,
        ["Overall", "Grand Mean",
         f"{likert_sum.overall_mean:.2f}", interp_fn(likert_sum.overall_mean)],
        total=True,
    )
    doc.add_paragraph()


# ══════════════════════════════════════════════════════════════
# Frequency tables
# ══════════════════════════════════════════════════════════════

def _write_frequency_table(doc: Document, ft: FrequencyTable,
                            table_num: int) -> None:
    """Write one frequency distribution table."""
    doc.add_paragraph()
    _caption(doc, f"Table {table_num}: Frequency Distribution of {ft.label}")
    doc.add_paragraph()

    n_rows = 1 + len(ft.rows) + 1
    table  = _make_table(doc, n_rows, 4)
    headers = [ft.label, "Frequency", "Percent (%)", "Cumulative (%)"]
    widths  = [6.5, 2.5, 2.5, 3.0]
    _header_row(table, headers, widths)

    for idx, row in enumerate(ft.rows):
        shade = (idx % 2 == 0)
        _data_row(
            table, idx + 1,
            [row.value, row.frequency,
             f"{row.percent:.1f}", f"{row.cumulative:.1f}"],
            shade=shade,
        )

    _data_row(
        table, len(ft.rows) + 1,
        ["Total", ft.n_valid, "100.0", ""],
        total=True,
    )
    doc.add_paragraph()


# ══════════════════════════════════════════════════════════════
# Crosstabulation tables
# ══════════════════════════════════════════════════════════════

def _write_crosstab_table(doc: Document, ct: CrosstabResult,
                           table_num: int) -> None:
    """Write one crosstabulation table with chi-square stats block."""
    doc.add_paragraph()
    _caption(doc,
             f"Table {table_num}: Crosstabulation of {ct.row_label} "
             f"and {ct.col_label}")
    doc.add_paragraph()

    n_cols  = 1 + len(ct.col_categories) + 1   # label + cols + total
    n_rows  = 1 + len(ct.row_categories) + 1   # header + rows + total
    table   = _make_table(doc, n_rows, n_cols)

    col_headers = [ct.row_label] + [str(c) for c in ct.col_categories] + ["Total"]
    widths      = [4.0] + [2.5] * len(ct.col_categories) + [1.8]
    _header_row(table, col_headers, widths)

    for idx, (row_cat, obs_row, row_tot) in enumerate(
            zip(ct.row_categories, ct.observed, ct.row_totals)):
        shade = (idx % 2 == 0)
        _data_row(
            table, idx + 1,
            [str(row_cat)] + list(obs_row) + [row_tot],
            shade=shade,
        )

    _data_row(
        table, len(ct.row_categories) + 1,
        ["Total"] + list(ct.col_totals) + [ct.n],
        total=True,
    )
    doc.add_paragraph()

    # Chi-square statistics block
    sig_text = "significant" if ct.significant else "not significant"
    p_str    = f"{ct.p_value:.4f}" if ct.p_value >= 0.0001 else "< 0.0001"
    _body(doc,
          f"χ²({ct.df}) = {ct.chi2:.3f}, p = {p_str}, "
          f"Cramer's V = {ct.cramers_v:.3f}  [{sig_text} at α = 0.05]",
          italic=True)
    if ct.note:
        _body(doc, f"Note: {ct.note}", italic=True)
    doc.add_paragraph()



# ══════════════════════════════════════════════════════════════
# Reliability table builders
# ══════════════════════════════════════════════════════════════

def _write_reliability_summary(doc: "Document", report: "ReliabilityReport",
                                table_num: int) -> int:
    """Summary table: one row per section + overall row."""
    import numpy as np
    doc.add_paragraph()
    _caption(doc, f"Table {table_num}: Cronbach\u2019s Alpha Reliability Coefficients")
    doc.add_paragraph()

    rows_data = report.to_summary_rows()
    n_rows    = 1 + len(rows_data)
    table     = _make_table(doc, n_rows, 6)
    headers   = ["Section", "Title", "Items", "N", "\u03b1", "Interpretation"]
    widths    = [1.8, 6.5, 1.2, 1.2, 1.5, 3.3]
    _header_row(table, headers, widths)

    for idx, row in enumerate(rows_data):
        is_total = (idx == len(rows_data) - 1)
        shade    = (idx % 2 == 0)
        _data_row(table, idx + 1, list(row), total=is_total, shade=shade)

    doc.add_paragraph()
    return table_num + 1


def _write_reliability_item_table(doc: "Document",
                                   sec_rel: "SectionReliability",
                                   table_num: int) -> int:
    """Per-section item-level reliability table."""
    import numpy as np
    doc.add_paragraph()
    _caption(doc,
             f"Table {table_num}: Item Analysis for Section {sec_rel.section_key} "
             f"\u2014 {sec_rel.section_title}")
    doc.add_paragraph()

    n_rows = 1 + len(sec_rel.items)
    table  = _make_table(doc, n_rows, 7)
    headers = ["Item No.", "Statement", "Mean", "SD",
               "r (item-total)", "r Interpretation", "\u03b1 if Deleted"]
    widths  = [1.5, 5.8, 1.2, 1.2, 1.8, 3.5, 1.8]
    _header_row(table, headers, widths)

    for idx, item in enumerate(sec_rel.items):
        shade = (idx % 2 == 0)
        _data_row(table, idx + 1, list(item.to_row()), shade=shade)

    doc.add_paragraph()
    return table_num + 1


# ══════════════════════════════════════════════════════════════
# Public API
# ══════════════════════════════════════════════════════════════

def export_word(
    dataset:             Dataset,
    questionnaire:       Questionnaire,
    variable_dictionary: VariableDictionary,
    validation_report:   ValidationReport,
    likert_sum:          LikertSummary,
    freq_tables:         list[FrequencyTable],
    crosstab_results:    list[CrosstabResult],
    output_dir:          str | Path,
    study_title:         str = "Research Study",
    seed:                int = 42,
    reliability_report:  "ReliabilityReport | None" = None,
) -> Path:
    """
    Export a submission-ready Chapter Four Word document.

    Parameters
    ----------
    dataset             : the generated Dataset
    questionnaire       : Questionnaire (for section structure and titles)
    variable_dictionary : VariableDictionary (for variable labels)
    validation_report   : ValidationReport (summary printed in cover block)
    likert_sum          : LikertSummary from analysis.descriptives.likert_summary()
    freq_tables         : list[FrequencyTable] from analysis.frequencies.all_categorical()
    crosstab_results    : list[CrosstabResult] from analysis.crosstabs.crosstab()
    reliability_report  : ReliabilityReport from analysis.reliability.full_reliability() — optional
    output_dir          : directory where the .docx file will be written
    study_title         : study title for the document heading
    seed                : random seed used (recorded in the document)

    Returns
    -------
    Path  — absolute path to the written .docx file
    """
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    stamp    = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe     = "".join(c if c.isalnum() or c in " _-" else "_"
                       for c in study_title)[:40].strip()
    filename = f"{safe.replace(' ', '_')}_{stamp}.docx"
    filepath = output_dir / filename

    doc = Document()
    _set_margins(doc)

    # ── Cover block ─────────────────────────────────────────
    doc.add_paragraph()
    _heading(doc, "CHAPTER FOUR", level=1)
    _heading(doc, "PRESENTATION AND ANALYSIS OF DATA", level=2)
    doc.add_paragraph()
    _body(doc, f"Study: {study_title}")
    _body(doc, f"Generated: {datetime.now().strftime('%B %d, %Y')}")
    _body(doc, f"Respondents: {len(dataset)}   |   Seed: {seed}")
    _body(doc, f"Validation: {validation_report.summary()}")
    doc.add_paragraph()

    # ── 4.1 Descriptive Statistics ───────────────────────────
    _heading(doc, "4.1 Descriptive Statistics", level=2)
    _body(doc,
          "The following tables present the mean scores and standard deviations "
          "for all Likert items, grouped by section.")
    doc.add_paragraph()

    table_num = 1
    for sec in questionnaire.sections:
        items        = likert_sum.items_for_section(sec.key)
        section_mean = likert_sum.section_means.get(sec.key, 0.0)
        if not items:
            continue
        _write_section_table(
            doc, sec.key, sec.title, items, section_mean, table_num
        )
        table_num += 1

    # Summary of section means
    _write_summary_table(doc, likert_sum, questionnaire, table_num)
    table_num += 1


    # ── 4.2 Reliability Analysis ────────────────────────────
    if reliability_report is not None:
        doc.add_page_break()
        _heading(doc, "4.2 Reliability Analysis (Cronbach\u2019s Alpha)", level=2)
        _body(doc,
              "Cronbach\u2019s alpha (\u03b1) was computed to assess the internal "
              "consistency of each subscale. An alpha coefficient of 0.70 or above "
              "is generally considered acceptable (Nunnally & Bernstein, 1994).")
        doc.add_paragraph()

        # Summary table
        table_num = _write_reliability_summary(doc, reliability_report, table_num)

        # Per-section item tables
        for sec_rel in reliability_report.sections:
            if sec_rel.items:
                table_num = _write_reliability_item_table(doc, sec_rel, table_num)

        doc.add_paragraph()

    # ── (renumber) Frequency Distributions ─────────────────────────
    if freq_tables:
        doc.add_page_break()
        _heading(doc, "4.3 Frequency Distributions", level=2)
        _body(doc,
              "The following tables present the frequency distributions "
              "of the demographic and categorical variables.")
        doc.add_paragraph()
        for ft in freq_tables:
            if ft.n_valid > 0:
                _write_frequency_table(doc, ft, table_num)
                table_num += 1

    # ── 4.3 Cross-Tabulations ───────────────────────────────
    if crosstab_results:
        doc.add_page_break()
        _heading(doc, "4.4 Cross-Tabulations", level=2)
        _body(doc,
              "The following tables examine the associations between demographic "
              "variables and satisfaction category using chi-square tests.")
        doc.add_paragraph()
        for ct in crosstab_results:
            _write_crosstab_table(doc, ct, table_num)
            table_num += 1

    doc.save(str(filepath))
    return filepath


# ══════════════════════════════════════════════════════════════
# Interpretation helper
# ══════════════════════════════════════════════════════════════

def _interpret_mean(mean: float) -> str:
    """Convert a Likert mean score to a verbal interpretation (5-point scale)."""
    if mean < 1.5:   return "Very Dissatisfied"
    if mean < 2.5:   return "Dissatisfied"
    if mean < 3.5:   return "Neutral"
    if mean < 4.5:   return "Satisfied"
    return "Very Satisfied"
