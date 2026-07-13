"""
research_engine/cli/project_commands.py

CLI commands for the Project / Writer mode.

Commands
--------
    project new     --title TITLE --level LEVEL [--guideline FILE]
    project upload  --session SID --file FILE
    project write   --session SID --chapter N|all [--model MODEL]
    project status  --session SID
    project export  --session SID [--format docx|txt|md]
    project dataset --session SID
    project list
"""
from __future__ import annotations

import os
import sys
import json
from pathlib import Path
from datetime import datetime


# ── Colour helpers (re-exported from parent) ──────────────────
def _c(text, code):
    if not sys.stdout.isatty(): return text
    return f"\033[{code}m{text}\033[0m"

green  = lambda t: _c(t, "32")
yellow = lambda t: _c(t, "33")
red    = lambda t: _c(t, "31")
gold   = lambda t: _c(t, "33;1")
bold   = lambda t: _c(t, "1")
cyan   = lambda t: _c(t, "36")
dim    = lambda t: _c(t, "2")


# ══════════════════════════════════════════════════════════════
# Session store helpers
# ══════════════════════════════════════════════════════════════

def _sessions_dir(project_root: Path) -> Path:
    d = project_root / "sessions"
    d.mkdir(exist_ok=True)
    return d


def _session_path(project_root: Path, session_id: str) -> Path:
    return _sessions_dir(project_root) / f"{session_id}.json"


def _load_session(project_root: Path, session_id: str):
    from research_engine.writer import ProjectSession
    path = _session_path(project_root, session_id)
    if not path.exists():
        print(red(f"  Session not found: {session_id}"))
        print(dim(f"  Expected at: {path}"))
        print(dim(f"  Run: python main.py project list"))
        return None
    return ProjectSession.from_file(path)


def _save_session(session, project_root: Path):
    path = _session_path(project_root, session.session_id)
    session.save(path)
    return path


def _list_sessions(project_root: Path) -> list[Path]:
    return sorted(_sessions_dir(project_root).glob("*.json"))


# ══════════════════════════════════════════════════════════════
# project new
# ══════════════════════════════════════════════════════════════

def cmd_project_new(args, project_root: Path) -> int:
    from research_engine.writer import (
        ProjectSession, extract_text, parse_guideline,
        extract_metadata_with_ai
    )

    title     = getattr(args, "title", "") or ""
    level     = getattr(args, "level", "") or ""
    guideline = getattr(args, "guideline", None)

    print(gold("\n  📋 Research Analysis Toolkit — New Project Session"))
    print()

    session = ProjectSession.new(title=title, level=level)

    # If a guideline file was provided, load + parse it
    if guideline:
        gpath = Path(guideline)
        if not gpath.exists():
            print(red(f"  File not found: {gpath}"))
            return 1
        print(f"  📄 Parsing guideline: {gpath.name}…")
        try:
            text = extract_text(gpath)
            session.guideline_raw = text
            session.add_file(gpath.name, text)
            meta = parse_guideline(text)
            # Merge: command-line title/level override the parsed values
            if title: meta.title = title
            if level:
                from research_engine.writer import EducationLevel
                meta.level = EducationLevel.detect(level)
            session.metadata = meta
            print(green(f"  ✅ Guideline parsed ({len(text.split())} words extracted)"))
        except Exception as exc:
            print(yellow(f"  ⚠️  Could not parse guideline: {exc}"))

    # Save
    path = _save_session(session, project_root)
    print()
    print(bold("  Session created:"))
    _print_session_summary(session)
    print()
    print(dim(f"  Saved: {path}"))
    print()
    print(f"  Next step:")
    print(cyan(f"    python main.py project write --session {session.session_id} --chapter 1"))
    return 0


# ══════════════════════════════════════════════════════════════
# project upload
# ══════════════════════════════════════════════════════════════

def cmd_project_upload(args, project_root: Path) -> int:
    from research_engine.writer import extract_text, parse_guideline

    session = _load_session(project_root, args.session)
    if session is None: return 1

    fpath = Path(args.file)
    if not fpath.exists():
        print(red(f"  File not found: {fpath}"))
        return 1

    print(f"  📄 Uploading: {fpath.name}…")
    try:
        text = extract_text(fpath)
        session.add_file(fpath.name, text)
        if fpath.suffix.lower() in (".docx", ".pdf", ".txt", ".md"):
            # Also treat as guideline if no guideline yet
            if not session.guideline_raw:
                session.guideline_raw = text
                meta = parse_guideline(text)
                # Preserve existing metadata
                for attr in ("title", "level", "institution", "department",
                             "objectives", "research_questions", "hypotheses"):
                    existing = getattr(session.metadata, attr)
                    parsed   = getattr(meta, attr)
                    if not existing and parsed:
                        setattr(session.metadata, attr, parsed)
                print(green(f"  ✅ File added as project guideline"))
            else:
                print(green(f"  ✅ File added to session"))
        _save_session(session, project_root)
        print(dim(f"  Files in session: {list(session.uploaded_files.keys())}"))
    except Exception as exc:
        print(red(f"  Failed: {exc}"))
        return 1

    return 0


# ══════════════════════════════════════════════════════════════
# project write
# ══════════════════════════════════════════════════════════════

def cmd_project_write(args, project_root: Path) -> int:
    from research_engine.writer import write_chapter, CHAPTER_TITLES

    session = _load_session(project_root, args.session)
    if session is None: return 1

    model   = getattr(args, "model", "gpt-4o") or "gpt-4o"
    chapter = str(getattr(args, "chapter", "1"))

    # Determine which chapters to write
    if chapter.lower() == "all":
        to_write = session.chapters_remaining or list(range(1, 6))
    else:
        try:
            to_write = [int(chapter)]
        except ValueError:
            print(red(f"  Invalid chapter: {chapter!r} (use 1–5 or 'all')"))
            return 1

    api_key = os.environ.get("OPENAI_API_KEY", "")
    if not api_key:
        print(red("  OPENAI_API_KEY environment variable not set."))
        print(dim("  Set it with: export OPENAI_API_KEY=sk-..."))
        return 1

    for n in to_write:
        title = CHAPTER_TITLES.get(n, f"Chapter {n}")
        print(gold(f"\n  ✍️  Writing Chapter {n}: {title}…"))
        wc_target = session.metadata.word_count_for_level().get(n, 2500)
        print(dim(f"     Target: ~{wc_target:,} words | Model: {model}"))

        try:
            result = write_chapter(session, n, api_key=api_key, model=model)
            _save_session(session, project_root)
            print(green(f"  ✅ Chapter {n} written ({result.word_count:,} words)"))
            print(dim(f"     {result.model_notes}"))
            print()
            print(cyan(f"  Preview:"))
            print(dim(f"  {result.preview(300)}"))
        except Exception as exc:
            print(red(f"  ✗ Failed to write Chapter {n}: {exc}"))
            if "api_key" in str(exc).lower() or "auth" in str(exc).lower():
                return 1
            # Continue with remaining chapters

    print()
    _print_progress(session)
    return 0


# ══════════════════════════════════════════════════════════════
# project status
# ══════════════════════════════════════════════════════════════

def cmd_project_status(args, project_root: Path) -> int:
    session = _load_session(project_root, args.session)
    if session is None: return 1

    print(gold(f"\n  📊 Session: {session.session_id}"))
    print()
    _print_session_summary(session)
    print()
    _print_progress(session)
    return 0


# ══════════════════════════════════════════════════════════════
# project export
# ══════════════════════════════════════════════════════════════

def cmd_project_export(args, project_root: Path) -> int:
    from research_engine.writer import CHAPTER_TITLES

    session = _load_session(project_root, args.session)
    if session is None: return 1

    fmt      = getattr(args, "format", "docx") or "docx"
    out_dir  = project_root / "output" / f"project_{session.session_id}"
    out_dir.mkdir(parents=True, exist_ok=True)

    print(gold(f"\n  📦 Exporting project ({fmt.upper()})…"))

    if fmt == "md":
        _export_markdown(session, out_dir)
    elif fmt == "txt":
        _export_txt(session, out_dir)
    else:
        _export_docx(session, out_dir)

    print(green(f"\n  ✅ Project exported to: {out_dir}"))
    return 0


def _export_markdown(session, out_dir: Path):
    from research_engine.writer import CHAPTER_TITLES
    for n in sorted(session.chapters.keys()):
        ch   = session.chapters[n]
        path = out_dir / f"chapter_{n:02d}_{ch.title.lower().replace(' ','_')}.md"
        path.write_text(f"# Chapter {n}: {ch.title}\n\n{ch.content}", encoding="utf-8")
        print(f"  ✅ {path.name}  ({ch.word_count:,} words)")


def _export_txt(session, out_dir: Path):
    from research_engine.writer import CHAPTER_TITLES
    full_path = out_dir / f"project_{session.session_id}_full.txt"
    parts = []
    for n in sorted(session.chapters.keys()):
        ch = session.chapters[n]
        parts.append(f"\n\n{'='*70}\nCHAPTER {n}: {ch.title.upper()}\n{'='*70}\n\n{ch.content}")
    full_path.write_text("\n".join(parts), encoding="utf-8")
    total_wc = sum(ch.word_count for ch in session.chapters.values())
    print(f"  ✅ {full_path.name}  ({total_wc:,} words total)")


def _export_docx(session, out_dir: Path):
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print(yellow("  python-docx not installed — exporting as Markdown instead"))
        _export_markdown(session, out_dir)
        return

    from research_engine.writer import CHAPTER_TITLES
    import re

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.54)

    # Title page
    m = session.metadata
    if m.title:
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(m.title.upper())
        run.bold     = True
        run.font.size = Pt(14)
    for field_val in [m.institution, m.department, m.student_name, m.year]:
        if field_val:
            p = doc.add_paragraph(field_val)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    for n in sorted(session.chapters.keys()):
        ch    = session.chapters[n]
        lines = ch.content.split("\n")

        for line in lines:
            line_s = line.strip()
            if not line_s:
                continue
            if line_s.startswith("### "):
                h = doc.add_heading(line_s[4:], level=3)
            elif line_s.startswith("## "):
                h = doc.add_heading(line_s[3:], level=2)
            elif line_s.startswith("# "):
                h = doc.add_heading(line_s[2:], level=1)
            else:
                p = doc.add_paragraph(line_s)
                p.style.font.size = Pt(12)

        doc.add_page_break()

    path = out_dir / f"project_{session.session_id}.docx"
    doc.save(str(path))
    total_wc = sum(ch.word_count for ch in session.chapters.values())
    print(f"  ✅ {path.name}  ({total_wc:,} words total)")


# ══════════════════════════════════════════════════════════════
# project dataset
# ══════════════════════════════════════════════════════════════

def cmd_project_dataset(args, project_root: Path) -> int:
    from research_engine.writer import suggest_study_config
    import json

    session = _load_session(project_root, args.session)
    if session is None: return 1

    print(gold(f"\n  🗄️  Generating study config from project metadata…"))
    config = suggest_study_config(session)

    study_name = f"project_{session.session_id}"
    study_dir  = project_root / "studies" / study_name
    study_dir.mkdir(parents=True, exist_ok=True)

    (study_dir / "config.json").write_text(
        json.dumps(config, indent=2), encoding="utf-8")

    print(green(f"  ✅ Study config created: studies/{study_name}/config.json"))
    print()
    print(dim("  To generate the dataset, run:"))
    print(cyan(f"    python main.py run --study {study_name}"))
    print()
    print(dim("  Then add questionnaire.json and demographics.json to the study directory"))
    print(dim("  and re-run for a full analysis-ready dataset."))
    return 0


# ══════════════════════════════════════════════════════════════
# project list
# ══════════════════════════════════════════════════════════════

def cmd_project_list(project_root: Path) -> int:
    from research_engine.writer import ProjectSession

    sessions = _list_sessions(project_root)
    if not sessions:
        print(yellow("  No project sessions found."))
        print(dim(f"  Start one with: python main.py project new --title '...' --level bsc"))
        return 0

    print(gold(f"\n  📋 Project Sessions ({len(sessions)} found)\n"))
    for path in sessions:
        try:
            s = ProjectSession.from_file(path)
            done  = s.chapters_done
            title = s.metadata.title or "(untitled)"
            level = s.metadata.level.value.upper()
            print(f"  {bold(s.session_id)}  {yellow(level):<8}  {title[:50]}")
            print(dim(f"             Chapters: {done if done else 'none'} | "
                      f"Updated: {s.updated_at[:10]}"))
        except Exception:
            print(dim(f"  {path.stem}  (could not read)"))
    print()
    return 0


# ══════════════════════════════════════════════════════════════
# Print helpers
# ══════════════════════════════════════════════════════════════

def _print_session_summary(session):
    from research_engine.writer import CHAPTER_TITLES
    m = session.metadata
    rows = [
        ("Session ID",  session.session_id),
        ("Title",       m.title or "(not set)"),
        ("Level",       m.level.value.upper()),
        ("Institution", m.institution or "(not set)"),
        ("Design",      m.research_design.value.replace("_"," ").title()),
        ("Citation",    m.citation_style),
        ("Population",  (m.population or "(not set)")[:60]),
        ("Objectives",  str(len(m.objectives))),
        ("Files",       ", ".join(session.uploaded_files.keys()) or "none"),
    ]
    for label, val in rows:
        print(f"  {dim(label+':'): <22} {val}")


def _print_progress(session):
    from research_engine.writer import CHAPTER_TITLES
    print(bold("  Chapter progress:"))
    for n in range(1, 6):
        title = CHAPTER_TITLES[n]
        ch    = session.chapters.get(n)
        if ch:
            wc_str = f"  {ch.word_count:,} words"
            print(f"    {green('✅')} Ch {n}: {title}{dim(wc_str)}")
        else:
            print(f"    {dim('⬜')} Ch {n}: {dim(title)}")

# ══════════════════════════════════════════════════════════════
# project revise  (Tier 1 #1)
# ══════════════════════════════════════════════════════════════

def cmd_project_revise(args, project_root: Path) -> int:
    from research_engine.writer import revise_chapter

    session = _load_session(project_root, args.session)
    if session is None: return 1

    chapter     = int(getattr(args, "chapter", 1))
    instruction = getattr(args, "instruction", "") or ""
    model       = getattr(args, "model", "gpt-4o") or "gpt-4o"

    if not instruction:
        print(red("  --instruction is required. Example:"))
        print(dim('  --instruction "Make the problem statement more focused"'))
        return 1

    existing = session.get_chapter(chapter)
    if existing is None:
        print(red(f"  Chapter {chapter} not written yet. Run 'project write' first."))
        return 1

    api_key = os.environ.get("OPENAI_API_KEY", "")
    if not api_key:
        print(red("  OPENAI_API_KEY not set.")); return 1

    print(gold(f"\n  ✏️  Revising Chapter {chapter}…"))
    print(dim(f"     Instruction: {instruction}"))
    print(dim(f"     Model: {model}"))

    try:
        result = revise_chapter(session, chapter, instruction, api_key=api_key, model=model)
        _save_session(session, project_root)
        print(green(f"  ✅ Chapter {chapter} revised ({result.word_count:,} words)"))
        print(dim(f"     {result.model_notes}"))
    except Exception as exc:
        print(red(f"  ✗ Revision failed: {exc}")); return 1
    return 0


# ══════════════════════════════════════════════════════════════
# project references  (Tier 1 #2)
# ══════════════════════════════════════════════════════════════

def cmd_project_references(args, project_root: Path) -> int:
    from research_engine.writer import generate_references, format_reference_list

    session = _load_session(project_root, args.session)
    if session is None: return 1

    fmt     = getattr(args, "format", "txt") or "txt"
    api_key = os.environ.get("OPENAI_API_KEY", "")

    print(gold("\n  📚 Generating reference list…"))
    refs = generate_references(session, api_key=api_key)
    print(green(f"  ✅ {len(refs)} unique references found"))

    out_dir = project_root / "output" / f"project_{session.session_id}"
    out_dir.mkdir(parents=True, exist_ok=True)

    if fmt == "md":
        text = refs.to_markdown()
        path = out_dir / "references.md"
    else:
        text = refs.to_text()
        path = out_dir / "references.txt"

    path.write_text(text, encoding="utf-8")
    print(dim(f"  Saved: {path}"))
    print()
    # Preview first 5
    for e in refs.entries[:5]:
        print(dim(f"  {e.full_text[:90]}"))
    if len(refs) > 5:
        print(dim(f"  ... and {len(refs)-5} more"))
    return 0


# ══════════════════════════════════════════════════════════════
# project build  (Tier 1 #3 — full study file generation)
# ══════════════════════════════════════════════════════════════

def cmd_project_build(args, project_root: Path) -> int:
    from research_engine.writer import save_study_files

    session = _load_session(project_root, args.session)
    if session is None: return 1

    api_key = os.environ.get("OPENAI_API_KEY", "")
    print(gold("\n  🔨 Building study config files (questionnaire + demographics)…"))

    try:
        created = save_study_files(session, project_root, api_key=api_key)
        _save_session(session, project_root)
        for fname, path in created.items():
            print(green(f"  ✅ {fname}  →  {path.relative_to(project_root)}"))
        study_name = f"project_{session.session_id}"
        print()
        print(dim("  To generate the dataset, run:"))
        print(cyan(f"    python main.py run --study {study_name}"))
    except Exception as exc:
        print(red(f"  ✗ Failed: {exc}")); return 1
    return 0


# ══════════════════════════════════════════════════════════════
# project ch4  (Tier 1 #4 — Ch4 with real data)
# ══════════════════════════════════════════════════════════════

def cmd_project_ch4(args, project_root: Path) -> int:
    from research_engine.writer import write_chapter4_with_data

    session = _load_session(project_root, args.session)
    if session is None: return 1

    api_key = os.environ.get("OPENAI_API_KEY", "")
    if not api_key:
        print(red("  OPENAI_API_KEY not set.")); return 1

    model = getattr(args, "model", "gpt-4o") or "gpt-4o"
    print(gold("\n  📊 Writing Chapter 4 with real dataset analysis…"))

    try:
        result = write_chapter4_with_data(session, project_root, api_key=api_key, model=model)
        _save_session(session, project_root)
        print(green(f"  ✅ Chapter 4 written ({result.word_count:,} words)"))
        print(dim(f"     {result.model_notes}"))
    except Exception as exc:
        print(red(f"  ✗ Failed: {exc}")); return 1
    return 0


# ══════════════════════════════════════════════════════════════
# project fullexport  (Tier 1 #5 — full submission docx)
# ══════════════════════════════════════════════════════════════

def cmd_project_fullexport(args, project_root: Path) -> int:
    from research_engine.exporters.project_docx_exporter import export_project_docx
    from research_engine.writer import generate_references

    session = _load_session(project_root, args.session)
    if session is None: return 1

    api_key  = os.environ.get("OPENAI_API_KEY", "")
    out_dir  = project_root / "output" / f"project_{session.session_id}"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / f"{session.session_id}_full_project.docx"

    print(gold("\n  📄 Generating full submission document…"))

    # Generate reference list first
    refs = None
    if session.chapters:
        print(dim("  Generating references…"))
        refs = generate_references(session, api_key=api_key)
        print(dim(f"  {len(refs)} references found"))

    try:
        result = export_project_docx(session, out_path, reference_list=refs)
        size_kb = result.stat().st_size // 1024
        print(green(f"  ✅ Full project exported: {result.name}  ({size_kb} KB)"))
        print(dim(f"  Location: {result}"))
        print()
        done = session.chapters_done
        total_wc = sum(session.chapters[n].word_count for n in done)
        print(dim(f"  Chapters included: {done}"))
        print(dim(f"  Total word count:  ~{total_wc:,} words"))
        print(dim(f"  References:        {len(refs) if refs else 0}"))
    except Exception as exc:
        print(red(f"  ✗ Export failed: {exc}")); import traceback; traceback.print_exc(); return 1
    return 0
